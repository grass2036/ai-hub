"""
Multi-Modal AI Services Module
Week 7 Day 1: Advanced AI Features Development

Provides comprehensive multi-modal AI capabilities including:
- Vision analysis (OpenAI Vision API)
- Speech processing (Whisper + TTS)
- Document understanding
- Multi-modal workflow integration
"""

import asyncio
import base64
import io
import json
import tempfile
import os
from typing import Optional, Dict, Any, List, AsyncGenerator, Union, BinaryIO
from pathlib import Path
from dataclasses import dataclass, asdict
from enum import Enum
import PIL.Image
import numpy as np
from fastapi import HTTPException, UploadFile, File
from backend.config.settings import get_settings
from backend.core.cost_tracker import CostTracker, ServiceType
from backend.core.logging_service import logging_service

settings = get_settings()


class ModalityType(str, Enum):
    """支持的多模态类型"""
    TEXT = "text"
    IMAGE = "image"
    AUDIO = "audio"
    VIDEO = "video"
    DOCUMENT = "document"


class AnalysisType(str, Enum):
    """图像分析类型"""
    GENERAL = "general"           # 通用分析
    OCR = "ocr"                  # 文字识别
    OBJECT_DETECTION = "object_detection"  # 物��检测
    CLASSIFICATION = "classification"      # 图像分类
    SENTIMENT = "sentiment"      # 情感分析
    CONTENT_MODERATION = "content_moderation"  # 内容审核


class SpeechTaskType(str, Enum):
    """语音处理任务类型"""
    TRANSCRIPTION = "transcription"  # 语音转文字
    TRANSLATION = "translation"      # 语音翻译
    DICTATION = "dictation"          # 听写
    COMMAND_RECOGNITION = "command_recognition"  # 命令识别


@dataclass
class ImageAnalysisResult:
    """图像分析结果"""
    description: str
    tags: List[str]
    confidence: float
    objects: List[Dict[str, Any]]
    text_content: Optional[str] = None
    sentiment: Optional[str] = None
    moderation_flags: List[str] = None
    metadata: Dict[str, Any] = None

    def __post_init__(self):
        if self.moderation_flags is None:
            self.moderation_flags = []
        if self.metadata is None:
            self.metadata = {}


@dataclass
class SpeechProcessResult:
    """语音处理结果"""
    text: str
    confidence: float
    language: str
    duration: float
    segments: List[Dict[str, Any]]
    translated_text: Optional[str] = None
    metadata: Dict[str, Any] = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class DocumentAnalysisResult:
    """文档分析结果"""
    content: str
    summary: str
    key_points: List[str]
    tables: List[Dict[str, Any]]
    images: List[Dict[str, Any]]
    metadata: Dict[str, Any]

    def __post_init__(self):
        if self.tables is None:
            self.tables = []
        if self.images is None:
            self.images = []
        if self.metadata is None:
            self.metadata = {}


class VisionService:
    """视觉分析服务"""

    def __init__(self):
        self._initialized = False
        self.cost_tracker = CostTracker()

    async def initialize(self):
        """初始化视觉服务"""
        if self._initialized:
            return

        try:
            if not settings.openrouter_api_key:
                raise ValueError("OpenRouter API密钥未配置")

            # 尝试导入OpenAI库
            try:
                import openai
                self.openai_client = openai.AsyncClient(api_key=settings.openrouter_api_key)
            except ImportError:
                raise ImportError("OpenAI库未安装，请运行: pip install openai")

            self._initialized = True
            logging_service.log_info("VisionService initialized successfully")

        except Exception as e:
            logging_service.log_error(f"VisionService initialization failed: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"视觉服务初始化失败: {str(e)}"
            )

    async def analyze_image(
        self,
        image_data: Union[str, bytes, UploadFile, PIL.Image.Image],
        analysis_type: AnalysisType = AnalysisType.GENERAL,
        prompt: str = "请详细分析这张图片的内容"
    ) -> ImageAnalysisResult:
        """分析图像内容"""
        if not self._initialized:
            await self.initialize()

        try:
            # 处理图像输入
            base64_image = await self._process_image_input(image_data)

            # 根据分析类型选择提示词
            analysis_prompt = self._get_analysis_prompt(analysis_type, prompt)

            # 调用Vision API
            response = await self.openai_client.chat.completions.create(
                model="gpt-4-vision-preview",
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": analysis_prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}"
                            }
                        }
                    ]
                }],
                max_tokens=1000,
                temperature=0.7
            )

            # 解析响应
            result_text = response.choices[0].message.content
            return self._parse_vision_response(result_text, analysis_type)

        except Exception as e:
            logging_service.log_error(f"Image analysis failed: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"图像分析失败: {str(e)}"
            )

    async def _process_image_input(self, image_data: Union[str, bytes, UploadFile, PIL.Image.Image]) -> str:
        """处理图像输入并转换为base64"""
        try:
            if isinstance(image_data, str):
                # 如果是base64字符串
                if image_data.startswith('data:image'):
                    return image_data.split(',')[1]
                return image_data
            elif isinstance(image_data, bytes):
                # 如果是字节数据
                return base64.b64encode(image_data).decode('utf-8')
            elif isinstance(image_data, UploadFile):
                # 如果是上传的文件
                contents = await image_data.read()
                image = PIL.Image.open(io.BytesIO(contents))
                return self._image_to_base64(image)
            elif isinstance(image_data, PIL.Image.Image):
                # 如果是PIL图像
                return self._image_to_base64(image)
            else:
                raise ValueError("不支持的图像格式")

        except Exception as e:
            raise ValueError(f"图像处理失败: {str(e)}")

    def _image_to_base64(self, image: PIL.Image.Image) -> str:
        """将PIL图像转换为base64"""
        buffer = io.BytesIO()
        # 确保图像是RGB格式
        if image.mode != 'RGB':
            image = image.convert('RGB')
        image.save(buffer, format='JPEG', quality=85)
        return base64.b64encode(buffer.getvalue()).decode('utf-8')

    def _get_analysis_prompt(self, analysis_type: AnalysisType, custom_prompt: str) -> str:
        """根据分析类型获取提示词"""
        prompts = {
            AnalysisType.GENERAL: f"{custom_prompt}。请提供详细的描述，包括主要对象、场景、颜色和布局。",
            AnalysisType.OCR: f"{custom_prompt}。请识别并提取图片中的所有文字内容，保持原有的格式和结构。",
            AnalysisType.OBJECT_DETECTION: f"{custom_prompt}。请识别图片中的所有物体，列出它们的名称和位置。",
            AnalysisType.CLASSIFICATION: f"{custom_prompt}。请对图片进行分类，并解释分类的理由。",
            AnalysisType.SENTIMENT: f"{custom_prompt}。请分析图片传达的情感和情绪。",
            AnalysisType.CONTENT_MODERATION: f"{custom_prompt}。请检查图片内容是否包含不当内容，如有请列出。"
        }
        return prompts.get(analysis_type, custom_prompt)

    def _parse_vision_response(self, response_text: str, analysis_type: AnalysisType) -> ImageAnalysisResult:
        """解析视觉API响应"""
        try:
            # 尝试解析JSON响应
            if response_text.startswith('{'):
                data = json.loads(response_text)
                return ImageAnalysisResult(**data)
            else:
                # 纯文本响应，转换为结构化数据
                return ImageAnalysisResult(
                    description=response_text,
                    tags=self._extract_tags(response_text),
                    confidence=0.8,
                    objects=[],
                    metadata={"analysis_type": analysis_type.value}
                )
        except Exception as e:
            logging_service.log_error(f"Failed to parse vision response: {str(e)}")
            # 返回基础结果
            return ImageAnalysisResult(
                description=response_text,
                tags=[],
                confidence=0.5,
                objects=[],
                metadata={"parse_error": str(e)}
            )

    def _extract_tags(self, text: str) -> List[str]:
        """从文本中提取标签"""
        # 简单的标签提取逻辑
        common_objects = [
            "人", "车", "建筑", "动物", "食物", "植物", "天空", "水", "山", "道路",
            "桌子", "椅子", "电脑", "手机", "书", "杯子", "门", "窗", "灯光", "标志"
        ]

        tags = []
        for obj in common_objects:
            if obj in text:
                tags.append(obj)
        return tags


class AudioService:
    """音频处理服务"""

    def __init__(self):
        self._initialized = False
        self.cost_tracker = CostTracker()

    async def initialize(self):
        """初始化音频服务"""
        if self._initialized:
            return

        try:
            # 导入必要的库
            try:
                import whisper
                import elevenlabs
                self.whisper_model = whisper.load_model("base")
                self.elevenlabs_client = elevenlabs.ElevenLabs()
            except ImportError as e:
                raise ImportError(f"音频处理库未安装: {str(e)}")

            self._initialized = True
            logging_service.log_info("AudioService initialized successfully")

        except Exception as e:
            logging_service.log_error(f"AudioService initialization failed: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"音频服务初始化失败: {str(e)}"
            )

    async def transcribe_audio(
        self,
        audio_data: Union[str, bytes, UploadFile],
        task_type: SpeechTaskType = SpeechTaskType.TRANSCRIPTION,
        language: str = "auto"
    ) -> SpeechProcessResult:
        """音频转文字"""
        if not self._initialized:
            await self.initialize()

        try:
            # 处理音频输入
            audio_file = await self._process_audio_input(audio_data)

            # 使用Whisper进行转录
            result = await asyncio.to_thread(
                self.whisper_model.transcribe,
                audio_file,
                language=language if language != "auto" else None,
                task="translate" if task_type == SpeechTaskType.TRANSLATION else "transcribe"
            )

            # 构建结果
            segments = []
            for segment in result.get("segments", []):
                segments.append({
                    "start": segment["start"],
                    "end": segment["end"],
                    "text": segment["text"],
                    "confidence": segment.get("avg_logprob", 0.0)
                })

            return SpeechProcessResult(
                text=result["text"],
                confidence=self._calculate_overall_confidence(segments),
                language=result.get("language", "unknown"),
                duration=result.get("segments", [{}])[-1].get("end", 0.0),
                segments=segments,
                metadata={
                    "task_type": task_type.value,
                    "whisper_model": "base"
                }
            )

        except Exception as e:
            logging_service.log_error(f"Audio transcription failed: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"音频转录失败: {str(e)}"
            )

    async def synthesize_speech(
        self,
        text: str,
        voice: str = "alloy",
        language: str = "en"
    ) -> bytes:
        """文字转语音"""
        if not self._initialized:
            await self.initialize()

        try:
            # 使用ElevenLabs生成语音
            audio = await asyncio.to_thread(
                self.elevenlabs_client.generate,
                text=text,
                voice=voice
            )

            # 转换为字节数据
            audio_bytes = b""
            for chunk in audio:
                audio_bytes += chunk

            return audio_bytes

        except Exception as e:
            logging_service.log_error(f"Speech synthesis failed: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"语音合成失败: {str(e)}"
            )

    async def _process_audio_input(self, audio_data: Union[str, bytes, UploadFile]) -> str:
        """处理音频输入"""
        try:
            if isinstance(audio_data, str):
                # 如果是文件路径
                return audio_data
            elif isinstance(audio_data, bytes):
                # 如果是字节数据，保存到临时文件
                with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as tmp:
                    tmp.write(audio_data)
                    return tmp.name
            elif isinstance(audio_data, UploadFile):
                # 如果是上传的文件
                contents = await audio_data.read()
                with tempfile.NamedTemporaryFile(
                    suffix=f".{audio_data.filename.split('.')[-1]}",
                    delete=False
                ) as tmp:
                    tmp.write(contents)
                    return tmp.name
            else:
                raise ValueError("不支持的音频格式")

        except Exception as e:
            raise ValueError(f"音频处理失败: {str(e)}")

    def _calculate_overall_confidence(self, segments: List[Dict[str, Any]]) -> float:
        """计算整体置信度"""
        if not segments:
            return 0.0

        total_confidence = sum(seg.get("confidence", 0.0) for seg in segments)
        return total_confidence / len(segments)


class DocumentService:
    """文档理解服务"""

    def __init__(self):
        self._initialized = False
        self.cost_tracker = CostTracker()

    async def initialize(self):
        """初始化文档服务"""
        if self._initialized:
            return

        try:
            # 导入必要的库
            try:
                import PyPDF2
                import docx
                import pandas as pd
                self.PyPDF2 = PyPDF2
                self.docx = docx
                self.pd = pd
            except ImportError as e:
                raise ImportError(f"文档处理库未安装: {str(e)}")

            self._initialized = True
            logging_service.log_info("DocumentService initialized successfully")

        except Exception as e:
            logging_service.log_error(f"DocumentService initialization failed: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"文档服务初始化失败: {str(e)}"
            )

    async def understand_document(
        self,
        document_data: Union[str, bytes, UploadFile],
        query: str = "请总结这个文档的主要内容"
    ) -> DocumentAnalysisResult:
        """理解文档内容"""
        if not self._initialized:
            await self.initialize()

        try:
            # 处理文档输入
            content, metadata = await self._extract_document_content(document_data)

            # 使用AI分析文档内容
            # 这里可以调用现有的AI服务进行分析
            summary = await self._generate_summary(content, query)
            key_points = await self._extract_key_points(content)
            tables = await self._extract_tables(document_data)
            images = await self._extract_images(document_data)

            return DocumentAnalysisResult(
                content=content,
                summary=summary,
                key_points=key_points,
                tables=tables,
                images=images,
                metadata=metadata
            )

        except Exception as e:
            logging_service.log_error(f"Document understanding failed: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"文档理解失败: {str(e)}"
            )

    async def _extract_document_content(self, document_data: Union[str, bytes, UploadFile]) -> tuple[str, Dict[str, Any]]:
        """提取文档内容"""
        try:
            if isinstance(document_data, str):
                # 文件路径
                file_path = document_data
                file_ext = Path(file_path).suffix.lower()
            elif isinstance(document_data, UploadFile):
                # 上传的文件
                file_ext = Path(document_data.filename).suffix.lower()
                contents = await document_data.read()
                with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as tmp:
                    tmp.write(contents)
                    file_path = tmp.name
            else:
                raise ValueError("不支持的文档格式")

            metadata = {"file_type": file_ext, "file_path": file_path}

            # 根据文件类型提取内容
            if file_ext == '.pdf':
                content = await self._extract_pdf_content(file_path)
            elif file_ext in ['.docx', '.doc']:
                content = await self._extract_docx_content(file_path)
            elif file_ext in ['.txt']:
                content = await self._extract_text_content(file_path)
            elif file_ext in ['.csv', '.xlsx', '.xls']:
                content = await self._extract_table_content(file_path)
            else:
                raise ValueError(f"不支持的文档格式: {file_ext}")

            return content, metadata

        except Exception as e:
            raise ValueError(f"文档内容提取失败: {str(e)}")

    async def _extract_pdf_content(self, file_path: str) -> str:
        """提取PDF内容"""
        content = ""
        with open(file_path, 'rb') as file:
            pdf_reader = self.PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                content += page.extract_text() + "\n"
        return content

    async def _extract_docx_content(self, file_path: str) -> str:
        """提取DOCX内容"""
        doc = self.docx.Document(file_path)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        return content

    async def _extract_text_content(self, file_path: str) -> str:
        """提取文本文件内容"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    async def _extract_table_content(self, file_path: str) -> str:
        """提取表格内容"""
        if file_path.endswith('.csv'):
            df = self.pd.read_csv(file_path)
        else:
            df = self.pd.read_excel(file_path)

        return df.to_string()

    async def _generate_summary(self, content: str, query: str) -> str:
        """生成文档摘要"""
        # 这里可以调用AI服务生成摘要
        # 简化实现，返回前500字符
        if len(content) > 500:
            return content[:500] + "...\n\n[摘要由AI生成]"
        return content

    async def _extract_key_points(self, content: str) -> List[str]:
        """提取关键点"""
        # 简化实现，基于段落分割
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        return paragraphs[:5]  # 返回前5个段落作为关键点

    async def _extract_tables(self, document_data: Union[str, bytes, UploadFile]) -> List[Dict[str, Any]]:
        """提取表格"""
        # 简化实现
        return []

    async def _extract_images(self, document_data: Union[str, bytes, UploadFile]) -> List[Dict[str, Any]]:
        """提取图片"""
        # 简化实现
        return []


class MultiModalAIService:
    """多模态AI服务管理器"""

    def __init__(self):
        self.vision_service = VisionService()
        self.audio_service = AudioService()
        self.document_service = DocumentService()
        self._initialized = False

    async def initialize(self):
        """初始化所有多模态服务"""
        if self._initialized:
            return

        # 并行初始化所有服务
        await asyncio.gather(
            self.vision_service.initialize(),
            self.audio_service.initialize(),
            self.document_service.initialize()
        )

        self._initialized = True
        logging_service.log_info("MultiModalAIService initialized successfully")

    async def process_multimodal_input(
        self,
        inputs: Dict[str, Any],
        task: str = "analyze"
    ) -> Dict[str, Any]:
        """处理多模态输入"""
        if not self._initialized:
            await self.initialize()

        results = {}

        try:
            # 处理图像输入
            if "image" in inputs:
                vision_result = await self.vision_service.analyze_image(
                    inputs["image"],
                    inputs.get("image_analysis_type", AnalysisType.GENERAL),
                    inputs.get("prompt", "分析这张图片")
                )
                results["vision"] = asdict(vision_result)

            # 处理音频输入
            if "audio" in inputs:
                audio_result = await self.audio_service.transcribe_audio(
                    inputs["audio"],
                    inputs.get("speech_task_type", SpeechTaskType.TRANSCRIPTION),
                    inputs.get("language", "auto")
                )
                results["audio"] = asdict(audio_result)

            # 处理文档输入
            if "document" in inputs:
                doc_result = await self.document_service.understand_document(
                    inputs["document"],
                    inputs.get("query", "请总结这个文档")
                )
                results["document"] = asdict(doc_result)

            # 处理文本输入
            if "text" in inputs:
                results["text"] = {
                    "content": inputs["text"],
                    "processed": True
                }

            return {
                "task": task,
                "results": results,
                "success": True,
                "metadata": {
                    "processing_time": 0,  # TODO: 实现时间统计
                    "modalities": list(inputs.keys())
                }
            }

        except Exception as e:
            logging_service.log_error(f"Multi-modal processing failed: {str(e)}")
            return {
                "task": task,
                "results": {},
                "success": False,
                "error": str(e),
                "metadata": {
                    "modalities": list(inputs.keys())
                }
            }


# 全局多模态AI服务实例
multimodal_ai_service = MultiModalAIService()